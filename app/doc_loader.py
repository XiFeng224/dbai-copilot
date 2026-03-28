from __future__ import annotations

import os
from typing import Iterable

from langchain_core.documents import Document


def _normalize_ws(s: str) -> str:
    # Collapse excessive whitespace to reduce embedding noise.
    return " ".join(s.split())


def load_file(file_path: str, file_name: str) -> list[Document]:
    """
    读取用户上传的文件，输出 LangChain Document 列表。
    由于此项目以 RAG demo 为主，采用「逐页/逐段」粗切分。
    """
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".pdf":
        return _load_pdf(file_path, file_name)
    if ext == ".docx":
        return _load_docx(file_path, file_name)
    if ext in [".txt", ".md", ".csv"]:
        return _load_text(file_path, file_name)
    raise ValueError(f"Unsupported file type: {ext}")


def _load_pdf(file_path: str, file_name: str) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    docs: list[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = _normalize_ws(text)
        if not text.strip():
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_name,
                    "page": i + 1,
                },
            )
        )
    return docs


def _load_docx(file_path: str, file_name: str) -> list[Document]:
    from docx import Document as DocxDocument

    d = DocxDocument(file_path)
    parts = []
    for p in d.paragraphs:
        t = _normalize_ws(p.text or "")
        if t.strip():
            parts.append(t)

    # Treat each paragraph as a document for better citations.
    docs: list[Document] = []
    for idx, text in enumerate(parts):
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_name,
                    "paragraph": idx + 1,
                },
            )
        )
    return docs


def _load_text(file_path: str, file_name: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    text = _normalize_ws(text)
    if not text.strip():
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_name,
                "chunk": 1,
            },
        )
    ]


def iter_metadata(doc: Document) -> str:
    src = doc.metadata.get("source", "unknown")
    page = doc.metadata.get("page")
    para = doc.metadata.get("paragraph")
    chunk = doc.metadata.get("chunk")

    if page:
        return f"{src} 第{page}页"
    if para:
        return f"{src} 段落{para}"
    if chunk:
        return f"{src} 段{chunk}"
    return f"{src}"

